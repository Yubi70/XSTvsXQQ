from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import pandas as pd
from docx import Document
from docx.shared import Inches


SRC = Path(__file__).parent
XST_HIST = SRC / "XST Historical Data (1).csv"
XQQ_HIST = SRC / "XQQ Historical Data (1).csv"

UP_THRESHOLD = 11.5
DOWN_THRESHOLD = 7.0

OUT_SWITCH_SIGNALS_5Y = SRC / "switch_signals_asym_last5y.png"
OUT_SWITCH_SIGNALS_2Y = SRC / "switch_signals_asym_last2y.png"
OUT_REAL_SWITCHES_5Y = SRC / "real_switches_asym_last5y.png"
OUT_REAL_SWITCHES_2Y = SRC / "real_switches_asym_last2y.png"
OUT_SWITCH_DURATION_5Y = SRC / "switch_duration_asym_last5y.png"
OUT_DOCX = SRC / "historical_asymmetric_theory.docx"

CSV_REAL_SWITCHES_5Y = SRC / "real_switches_asym_last5y.csv"
CSV_REAL_SWITCHES_2Y = SRC / "real_switches_asym_last2y.csv"


def load_hist(path: Path) -> pd.DataFrame:
    df = pd.read_csv(path, usecols=["Date", "Price"])
    df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    df["Price"] = pd.to_numeric(df["Price"].astype(str).str.replace(",", ""), errors="coerce")
    return df.dropna().sort_values("Date").reset_index(drop=True)


def load_merged() -> pd.DataFrame:
    xst = load_hist(XST_HIST).rename(columns={"Price": "Price_XST"})
    xqq = load_hist(XQQ_HIST).rename(columns={"Price": "Price_XQQ"})
    merged = pd.merge(xst, xqq, on="Date", how="inner")
    avg = (merged["Price_XST"] + merged["Price_XQQ"]) / 2
    merged["Delta %"] = ((merged["Price_XST"] - merged["Price_XQQ"]) / avg) * 100
    return merged.sort_values("Date").reset_index(drop=True)


def windowed(df: pd.DataFrame, years: int) -> pd.DataFrame:
    cutoff = df["Date"].max() - pd.DateOffset(years=years)
    return df[df["Date"] >= cutoff].copy().reset_index(drop=True)


def pct_change_since_start(series: pd.Series) -> float:
    if len(series) < 2 or series.iloc[0] == 0:
        return 0.0
    return (series.iloc[-1] / series.iloc[0] - 1.0) * 100.0


def change_caption(df: pd.DataFrame) -> str:
    xst_chg = pct_change_since_start(df["Price_XST"])
    xqq_chg = pct_change_since_start(df["Price_XQQ"])
    start = df["Date"].iloc[0].date()
    return f"Change since {start}: XST {xst_chg:+.2f}% | XQQ {xqq_chg:+.2f}%"


def compute_switches_asym(df: pd.DataFrame) -> pd.DataFrame:
    holding = "XST"
    rows = []
    switch_no = 0

    for _, row in df.iterrows():
        d = float(row["Delta %"])
        if d >= UP_THRESHOLD and holding == "XST":
            switch_no += 1
            rows.append(
                {
                    "Switch #": switch_no,
                    "Date": row["Date"],
                    "From": "XST",
                    "To": "XQQ",
                    "Signed premium %": round(d, 3),
                    "Price_XST": round(float(row["Price_XST"]), 2),
                    "Price_XQQ": round(float(row["Price_XQQ"]), 2),
                }
            )
            holding = "XQQ"
        elif d <= -DOWN_THRESHOLD and holding == "XQQ":
            switch_no += 1
            rows.append(
                {
                    "Switch #": switch_no,
                    "Date": row["Date"],
                    "From": "XQQ",
                    "To": "XST",
                    "Signed premium %": round(d, 3),
                    "Price_XST": round(float(row["Price_XST"]), 2),
                    "Price_XQQ": round(float(row["Price_XQQ"]), 2),
                }
            )
            holding = "XST"

    out = pd.DataFrame(
        rows,
        columns=["Switch #", "Date", "From", "To", "Signed premium %", "Price_XST", "Price_XQQ"],
    )
    return out


def plot_switch_signals_asym(df: pd.DataFrame, years: int, out_path: Path) -> None:
    switches = compute_switches_asym(df)
    up_sw = switches[switches["To"] == "XQQ"]
    down_sw = switches[switches["To"] == "XST"]

    fig, ax = plt.subplots(figsize=(14, 6))
    ax.plot(df["Date"], df["Delta %"], color="#1f77b4", linewidth=1.4, label="Delta %")
    ax.axhline(0, color="gray", linestyle=":", linewidth=1.0)
    ax.axhline(UP_THRESHOLD, color="#d32f2f", linestyle="--", linewidth=1.1, label=f"+{UP_THRESHOLD}%: XST->XQQ")
    ax.axhline(-DOWN_THRESHOLD, color="#2e7d32", linestyle="--", linewidth=1.1, label=f"-{DOWN_THRESHOLD}%: XQQ->XST")

    if not up_sw.empty:
        ax.scatter(up_sw["Date"], up_sw["Signed premium %"], s=70, c="#d32f2f", marker="v", label="Switch to XQQ", zorder=5)
    if not down_sw.empty:
        ax.scatter(down_sw["Date"], down_sw["Signed premium %"], s=70, c="#2e7d32", marker="^", label="Switch to XST", zorder=5)

    ax.set_title(
        f"XST vs XQQ Asymmetric Switch Signals - Last {years} Years\n"
        f"{change_caption(df)}",
        fontsize=12,
        fontweight="bold",
    )
    ax.set_ylabel("Delta % (symmetric)")
    ax.set_xlabel("Date")
    ax.grid(axis="y", linestyle=":", alpha=0.4)
    ax.legend(loc="upper left", fontsize=9)
    ax.spines[["top", "right"]].set_visible(False)

    fig.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    print(f"Saved {out_path}")


def plot_real_switches_asym(df_switches: pd.DataFrame, years: int, out_path: Path, caption_source: pd.DataFrame) -> None:
    fig, ax = plt.subplots(figsize=(14, 6))

    x = range(1, len(df_switches) + 1)
    y = pd.to_numeric(df_switches["Signed premium %"], errors="coerce")
    colors = ["#d32f2f" if t == "XQQ" else "#2e7d32" for t in df_switches["To"]]

    bars = ax.bar(x, y, color=colors, alpha=0.9)
    ax.axhline(0, color="gray", linestyle=":", linewidth=1)
    ax.set_ylabel("Signed premium %")
    ax.set_xlabel("Switch #")
    ax.set_title(
        f"Asymmetric Real Switch Premiums - Last {years} Years\n"
        f"{change_caption(caption_source)}",
        fontsize=12,
        fontweight="bold",
    )
    ax.grid(axis="y", linestyle=":", alpha=0.35)
    ax.spines[["top", "right"]].set_visible(False)

    for bar, val in zip(bars, y):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + (0.20 if val >= 0 else -0.20),
            f"{val:+.2f}%",
            ha="center",
            va="bottom" if val >= 0 else "top",
            fontsize=8,
        )

    ax.set_xticks(list(x))
    ax.set_xticklabels([str(i) for i in x])

    legend_handles = [
        mpatches.Patch(color="#d32f2f", label=f"Switch to XQQ (delta >= +{UP_THRESHOLD}%)"),
        mpatches.Patch(color="#2e7d32", label=f"Switch to XST (delta <= -{DOWN_THRESHOLD}%)"),
    ]
    ax.legend(handles=legend_handles, fontsize=9, loc="upper left")

    fig.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    print(f"Saved {out_path}")


def plot_switch_duration_asym(df_switches_5y: pd.DataFrame, caption_source: pd.DataFrame) -> None:
    if len(df_switches_5y) < 2:
        return

    durations = []
    ordered = df_switches_5y.sort_values("Date").reset_index(drop=True)

    for i in range(len(ordered) - 1):
        days = (ordered.loc[i + 1, "Date"] - ordered.loc[i, "Date"]).days
        label = f"#{ordered.loc[i, 'Switch #']}->#{ordered.loc[i + 1, 'Switch #']}\n{ordered.loc[i, 'Date'].strftime('%Y-%m-%d')}"
        durations.append({"label": label, "days": days, "from": ordered.loc[i, "From"]})

    dur_df = pd.DataFrame(durations)
    avg_days = dur_df["days"].mean()
    avg_xst_days = dur_df.loc[dur_df["from"] == "XST", "days"].mean()
    avg_xqq_days = dur_df.loc[dur_df["from"] == "XQQ", "days"].mean()

    c_xst = "#1f77b4"
    c_xqq = "#ff7f0e"
    colors = [c_xst if row["from"] == "XST" else c_xqq for _, row in dur_df.iterrows()]

    fig, ax = plt.subplots(figsize=(14, 6))
    bars = ax.bar(range(len(dur_df)), dur_df["days"], color=colors, edgecolor="white", linewidth=0.8)

    for bar, days in zip(bars, dur_df["days"]):
        x = bar.get_x() + bar.get_width() / 2
        if days >= 40:
            ax.text(x, days / 2, f"{days}d", ha="center", va="center", fontsize=9, fontweight="bold", color="white")
        else:
            ax.text(x, days + 4, f"{days}d", ha="center", va="bottom", fontsize=9, fontweight="bold", color="#333333")

    ax.axhline(avg_days, color="crimson", linestyle="--", linewidth=1.6, label=f"Overall avg {avg_days:.0f}d")
    if not pd.isna(avg_xst_days):
        ax.axhline(avg_xst_days, color=c_xst, linestyle=":", linewidth=1.3, label=f"XST avg {avg_xst_days:.0f}d")
    if not pd.isna(avg_xqq_days):
        ax.axhline(avg_xqq_days, color=c_xqq, linestyle=":", linewidth=1.3, label=f"XQQ avg {avg_xqq_days:.0f}d")

    ax.set_xticks(range(len(dur_df)))
    ax.set_xticklabels(dur_df["label"], fontsize=8.5)
    ax.set_ylabel("Holding duration (calendar days)")
    ax.set_title(
        "Asymmetric Duration Between Switches (last 5 years)\n"
        f"{change_caption(caption_source)}",
        fontsize=12,
        fontweight="bold",
    )
    ax.grid(axis="y", linestyle=":", alpha=0.45)
    ax.spines[["top", "right"]].set_visible(False)

    legend_handles = [
        mpatches.Patch(color=c_xst, label="Holding XST"),
        mpatches.Patch(color=c_xqq, label="Holding XQQ"),
    ]
    ax.legend(
        handles=legend_handles + [plt.Line2D([0], [0], color="crimson", linestyle="--", label=f"Overall avg {avg_days:.0f}d")],
        fontsize=9,
        loc="upper left",
    )

    fig.tight_layout()
    fig.savefig(OUT_SWITCH_DURATION_5Y, dpi=150, bbox_inches="tight")
    plt.close(fig)
    print(f"Saved {OUT_SWITCH_DURATION_5Y}")


def strategy_return(df: pd.DataFrame) -> tuple[float, int]:
    start_xst = float(df["Price_XST"].iloc[0])
    end_xst = float(df["Price_XST"].iloc[-1])
    end_xqq = float(df["Price_XQQ"].iloc[-1])

    holding = "XST"
    units = 1.0 / start_xst
    switches = 0

    for _, row in df.iterrows():
        d = float(row["Delta %"])
        pxst = float(row["Price_XST"])
        pxqq = float(row["Price_XQQ"])

        if d >= UP_THRESHOLD and holding == "XST":
            units = (units * pxst) / pxqq
            holding = "XQQ"
            switches += 1
        elif d <= -DOWN_THRESHOLD and holding == "XQQ":
            units = (units * pxqq) / pxst
            holding = "XST"
            switches += 1

    final_value = units * (end_xst if holding == "XST" else end_xqq)
    return (final_value - 1.0) * 100.0, switches


def passive_5050_return(df: pd.DataFrame) -> float:
    start_xst = float(df["Price_XST"].iloc[0])
    start_xqq = float(df["Price_XQQ"].iloc[0])
    end_xst = float(df["Price_XST"].iloc[-1])
    end_xqq = float(df["Price_XQQ"].iloc[-1])
    final_value = 0.5 * (end_xst / start_xst) + 0.5 * (end_xqq / start_xqq)
    return (final_value - 1.0) * 100.0


def add_metrics_table(doc: Document, data_rows: list[dict]) -> None:
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text = "Window"
    hdr[1].text = "Thresholds"
    hdr[2].text = "Real switches"
    hdr[3].text = "Strategy return %"
    hdr[4].text = "50/50 return %"
    hdr[5].text = "Diff vs 50/50 (pp)"

    for row in data_rows:
        cells = table.add_row().cells
        cells[0].text = row["window"]
        cells[1].text = row["thresholds"]
        cells[2].text = str(row["switches"])
        cells[3].text = f"{row['strategy']:.2f}%"
        cells[4].text = f"{row['passive']:.2f}%"
        cells[5].text = f"{row['edge']:+.2f}"


def build_docx(df2: pd.DataFrame, df5: pd.DataFrame) -> None:
    doc = Document()
    doc.add_heading("Historical Asymmetric Theory", level=1)
    doc.add_paragraph(
        "Directional thresholds: +11.5% for XST->XQQ and -7.0% for XQQ->XST, "
        "using symmetric delta ((XST - XQQ) / avg) * 100."
    )

    rows = []
    for label, frame in (("Last 2 years", df2), ("Last 5 years", df5)):
        strategy_ret, switches = strategy_return(frame)
        passive_ret = passive_5050_return(frame)
        rows.append(
            {
                "window": label,
                "thresholds": "+11.5% / -7.0%",
                "switches": switches,
                "strategy": strategy_ret,
                "passive": passive_ret,
                "edge": strategy_ret - passive_ret,
            }
        )

    doc.add_heading("Performance Summary", level=2)
    add_metrics_table(doc, rows)

    doc.add_heading("Charts", level=2)
    for title, img in [
        ("Asymmetric switch signals - last 5 years", OUT_SWITCH_SIGNALS_5Y),
        ("Asymmetric switch signals - last 2 years", OUT_SWITCH_SIGNALS_2Y),
        ("Asymmetric real switches - last 5 years", OUT_REAL_SWITCHES_5Y),
        ("Asymmetric real switches - last 2 years", OUT_REAL_SWITCHES_2Y),
        ("Asymmetric switch duration - last 5 years", OUT_SWITCH_DURATION_5Y),
    ]:
        doc.add_paragraph(title)
        if img.exists():
            doc.add_picture(str(img), width=Inches(6.8))

    doc.save(OUT_DOCX)
    print(f"Saved {OUT_DOCX}")


def main() -> None:
    merged = load_merged()
    df5 = windowed(merged, 5)
    df2 = windowed(merged, 2)

    sw5 = compute_switches_asym(df5)
    sw2 = compute_switches_asym(df2)

    sw5_out = sw5.copy()
    sw2_out = sw2.copy()
    sw5_out["Date"] = sw5_out["Date"].dt.strftime("%Y-%m-%d")
    sw2_out["Date"] = sw2_out["Date"].dt.strftime("%Y-%m-%d")
    sw5_out.to_csv(CSV_REAL_SWITCHES_5Y, index=False)
    sw2_out.to_csv(CSV_REAL_SWITCHES_2Y, index=False)
    print(f"Saved {CSV_REAL_SWITCHES_5Y}")
    print(f"Saved {CSV_REAL_SWITCHES_2Y}")

    plot_switch_signals_asym(df5, 5, OUT_SWITCH_SIGNALS_5Y)
    plot_switch_signals_asym(df2, 2, OUT_SWITCH_SIGNALS_2Y)
    plot_real_switches_asym(sw5, 5, OUT_REAL_SWITCHES_5Y, df5)
    plot_real_switches_asym(sw2, 2, OUT_REAL_SWITCHES_2Y, df2)
    plot_switch_duration_asym(sw5, df5)

    build_docx(df2, df5)


if __name__ == "__main__":
    main()
